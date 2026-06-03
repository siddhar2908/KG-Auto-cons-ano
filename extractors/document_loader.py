"""
extractors/document_loader.py
------------------------------
Loads and extracts text from:
  - PDF  (digital-native via pdfplumber)
  - PDF  (scanned / image-heavy via Surya OCR v0.17.x)
  - DOCX (via python-docx)
  - TXT  (plain read)

Also detects which pages are scanned vs text-native and identifies
pages that likely contain tables (for the table extractor).

Output per document:
    DocumentContent:
        pages: list[PageContent]   — one entry per page
        raw_text: str              — full concatenated text
        table_page_nums: list[int] — 0-indexed pages that likely have tables
        is_scanned: bool
"""

import re
import threading
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger
from PIL import Image

# ─── Surya OCR ────────────────────────────────────────────────────────────────
# Disabled by default — Surya v0.17.x has API instability issues.
# For digital PDFs (scanned=False), pdfplumber handles extraction perfectly.
# Re-enable by setting SURYA_ENABLED = True once you have scanned documents.
SURYA_ENABLED    = False   # ← set True to enable Surya for scanned PDFs
SURYA_AVAILABLE  = False
_surya_rec_predictor = None
_surya_det_predictor = None


# ─── Data structures ──────────────────────────────────────────────────────────

@dataclass
class PageContent:
    page_num: int       # 0-indexed
    text: str
    is_scanned: bool = False
    has_table: bool = False
    image_count: int = 0
    char_count: int = 0
    rect_count: int = 0         # number of rectangles (cell borders) on page
    has_grid_lines: bool = False # true if pdfplumber detected line-based grid

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class DocumentContent:
    doc_id: str
    source_path: Path
    doc_type: str          # "pdf" | "docx" | "txt"
    pages: list[PageContent] = field(default_factory=list)
    raw_text: str = ""
    table_page_nums: list[int] = field(default_factory=list)
    is_scanned: bool = False
    total_pages: int = 0
    sector: Optional[str] = None
    sector_confidence: float = 0.0


# ─── Surya OCR singleton ──────────────────────────────────────────────────────

_surya_load_lock = threading.Lock()


def _get_surya_predictors():
    """Returns Surya predictors if enabled and available, else (None, None)."""
    global _surya_rec_predictor, _surya_det_predictor, SURYA_AVAILABLE
    if not SURYA_ENABLED:
        return None, None
    if _surya_rec_predictor is not None:
        return _surya_det_predictor, _surya_rec_predictor
    with _surya_load_lock:
        if _surya_rec_predictor is not None:
            return _surya_det_predictor, _surya_rec_predictor
        logger.info("Loading Surya OCR models...")
        try:
            from surya.recognition import RecognitionPredictor
            from surya.detection import DetectionPredictor
            _surya_det_predictor = DetectionPredictor()
            try:
                _surya_rec_predictor = RecognitionPredictor(
                    foundation_predictor=_surya_det_predictor
                )
            except TypeError:
                _surya_rec_predictor = RecognitionPredictor()
            SURYA_AVAILABLE = True
            logger.success("Surya OCR models loaded.")
        except Exception as e:
            logger.error(f"Surya model load failed: {e}")
            return None, None
    return _surya_det_predictor, _surya_rec_predictor


def _surya_ocr_page(image: Image.Image) -> str:
    """Run Surya OCR v0.17.x on a PIL Image. Returns extracted text."""
    det_predictor, rec_predictor = _get_surya_predictors()
    if det_predictor is None or rec_predictor is None:
        return ""
    try:
        # v0.17.x API: predictors are callable
        det_result = det_predictor([image])
        bboxes = [r.bboxes for r in det_result]
        rec_result = rec_predictor([image], bboxes)
        lines = []
        for page_result in rec_result:
            for line in page_result.text_lines:
                if line.text.strip():
                    lines.append(line.text)
        return "\n".join(lines)
    except Exception as e:
        logger.debug(f"Surya OCR page failed: {e}")
        return ""


# ─── Page analysis helpers ────────────────────────────────────────────────────

_TABLE_KEYWORDS = re.compile(
    r"\b(table|schedule|statement|bill of quantities|boq|summary|list of|rates?)\b",
    re.IGNORECASE
)
_MIN_TEXT_CHARS_FOR_NATIVE = 80  # fewer chars → likely scanned


def _page_likely_has_table(page: pdfplumber.page.Page, text: str) -> bool:
    """Heuristic: does this page likely contain a table?"""
    if _TABLE_KEYWORDS.search(text):
        return True
    # Check for grid lines (many horizontal/vertical lines = table)
    if len(page.lines) > 10:
        return True
    # Check for many rects (cells)
    if len(page.rects) > 6:
        return True
    return False


# ─── PDF loader ───────────────────────────────────────────────────────────────

def _load_pdf(path: Path, doc_id: str) -> DocumentContent:
    doc = DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="pdf",
    )

    with pdfplumber.open(path) as pdf:
        doc.total_pages = len(pdf.pages)
        logger.info(f"Loading PDF: {path.name} ({doc.total_pages} pages)")

        # Determine if whole PDF is scanned by sampling first 5 pages
        sample_chars = sum(
            len(pdf.pages[i].extract_text() or "")
            for i in range(min(5, doc.total_pages))
        )
        doc.is_scanned = (sample_chars < _MIN_TEXT_CHARS_FOR_NATIVE * 5)
        if doc.is_scanned:
            logger.info("Document appears to be scanned — will use Surya OCR.")

        for i, page in enumerate(pdf.pages):
            native_text = page.extract_text() or ""
            is_scanned_page = len(native_text.strip()) < _MIN_TEXT_CHARS_FOR_NATIVE

            if is_scanned_page and SURYA_AVAILABLE:
                # Rasterise page for Surya OCR
                pil_img = page.to_image(resolution=200).original
                page_text = _surya_ocr_page(pil_img)
                # If surya returned nothing, fall back to native
                if not page_text.strip():
                    page_text = native_text
            else:
                page_text = native_text

            has_table = _page_likely_has_table(page, page_text)
            image_count = len(page.images)
            rect_count = len(page.rects)
            has_grid_lines = len(page.lines) > 8

            pc = PageContent(
                page_num=i,
                text=page_text,
                is_scanned=is_scanned_page,
                has_table=has_table,
                image_count=image_count,
                rect_count=rect_count,
                has_grid_lines=has_grid_lines,
            )
            doc.pages.append(pc)
            if has_table:
                doc.table_page_nums.append(i)

    doc.raw_text = "\n\n".join(p.text for p in doc.pages if p.text)
    logger.success(
        f"PDF loaded: {doc.total_pages} pages, "
        f"{len(doc.table_page_nums)} table pages, "
        f"scanned={doc.is_scanned}"
    )
    return doc


# ─── DOCX loader ──────────────────────────────────────────────────────────────

def _load_docx(path: Path, doc_id: str) -> DocumentContent:
    doc = DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="docx",
    )
    docx = DocxDocument(path)
    full_text_parts = []
    table_count = 0

    for para in docx.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)

    for i, table in enumerate(docx.tables):
        table_count += 1
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        full_text_parts.append(f"\n[TABLE {i+1}]\n{table_text}\n[/TABLE]\n")

    full_text = "\n".join(full_text_parts)
    pc = PageContent(page_num=0, text=full_text, has_table=(table_count > 0))
    doc.pages = [pc]
    doc.raw_text = full_text
    doc.total_pages = 1
    if table_count > 0:
        doc.table_page_nums = [0]

    logger.success(f"DOCX loaded: {len(docx.paragraphs)} paragraphs, {table_count} tables")
    return doc


# ─── TXT loader ───────────────────────────────────────────────────────────────

def _load_txt(path: Path, doc_id: str) -> DocumentContent:
    text = path.read_text(encoding="utf-8", errors="replace")
    pc = PageContent(page_num=0, text=text)
    return DocumentContent(
        doc_id=doc_id,
        source_path=path,
        doc_type="txt",
        pages=[pc],
        raw_text=text,
        total_pages=1,
    )


# ─── Public API ───────────────────────────────────────────────────────────────

def load_document(path: Path, doc_id: str = None) -> DocumentContent:
    """
    Load any supported document and return a DocumentContent object.
    Automatically chooses the right loader based on file extension.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    _id = doc_id or path.stem
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(path, _id)
    elif suffix in (".docx", ".doc"):
        return _load_docx(path, _id)
    elif suffix == ".txt":
        return _load_txt(path, _id)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")